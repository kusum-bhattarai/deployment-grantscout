"""
Smart Ingest Pipeline (The Librarian)
======================================
Converts raw documents into indexed, citation-ready knowledge chunks.

Key Features:
- PDF/DOCX parsing with page-level tracking
- Metadata tagging for every chunk (source_id, page, section)
- FAISS vector store for semantic retrieval
- Citation-ready output
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import fitz  # PyMuPDF
from docx import Document
import numpy as np
from PIL import Image
import io
import pytesseract
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document as LangchainDocument
from dotenv import load_dotenv

load_dotenv()


class SmartChunker:
    """
    Intelligent document chunker that preserves citation metadata.
    """

    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None
    ):
        self.chunk_size = chunk_size or int(os.getenv("CHUNK_SIZE", 1000))
        self.chunk_overlap = chunk_overlap or int(os.getenv("CHUNK_OVERLAP", 200))

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def parse_pdf(self, file_path: str, doc_type: str = "solicitation") -> List[LangchainDocument]:
        """
        Parse PDF and extract text with page-level metadata.

        Args:
            file_path: Path to PDF file
            doc_type: Document type (solicitation, narrative, etc.)

        Returns:
            List of LangchainDocument objects with metadata
        """
        documents = []
        file_name = Path(file_path).name

        print(f"📄 Parsing PDF: {file_name}")

        try:
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc, 1):
                    # IMPROVED: Use layout-preserving text extraction
                    # This better preserves table structure and formatting
                    text = page.get_text("text", sort=True)

                    # Also try to extract tables specifically
                    tables_text = self._extract_tables_from_page(page)
                    if tables_text:
                        text += "\n\n" + tables_text

                    # NEW: Extract and OCR images (for table images)
                    ocr_text = self._extract_and_ocr_images(page, page_num)
                    if ocr_text:
                        text += "\n\n[OCR from images on this page:]\n" + ocr_text

                    # Skip empty pages
                    if not text.strip():
                        continue

                    # Detect section headers (common patterns)
                    section = self._detect_section(text)

                    # Create chunks from this page
                    chunks = self.text_splitter.split_text(text)

                    for chunk_idx, chunk in enumerate(chunks):
                        metadata = {
                            "source": file_name,
                            "source_type": doc_type,
                            "page": page_num,
                            "section": section,
                            "chunk_index": chunk_idx,
                            "total_chunks_on_page": len(chunks)
                        }

                        documents.append(LangchainDocument(
                            page_content=chunk,
                            metadata=metadata
                        ))

                print(f"   ✅ Extracted {len(documents)} chunks from {page_num} pages")

        except Exception as e:
            print(f"   ❌ Error parsing PDF: {e}")

        return documents

    def parse_docx(self, file_path: str, doc_type: str = "narrative") -> List[LangchainDocument]:
        """
        Parse DOCX and extract text with paragraph-level metadata.
        NOW INCLUDES: Native table extraction + OCR for embedded images

        Args:
            file_path: Path to DOCX file
            doc_type: Document type

        Returns:
            List of LangchainDocument objects with metadata
        """
        documents = []
        file_name = Path(file_path).name

        print(f"📝 Parsing DOCX: {file_name}")

        try:
            doc = Document(file_path)
            current_section = "Introduction"
            page_estimate = 1  # Approximate page number (rough estimation)
            char_count = 0

            full_text = ""
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect section headers
                if self._is_header(para):
                    current_section = text
                    print(f"   📍 Section detected: {current_section}")

                full_text += text + "\n\n"
                char_count += len(text)

                # Rough page estimation (2500 chars ≈ 1 page)
                if char_count > 2500:
                    page_estimate += 1
                    char_count = 0

            # IMPORTANT: Extract native DOCX tables (if any)
            table_text = self._extract_tables_from_docx(doc)
            if table_text:
                print(f"   ✅ Extracted {len(doc.tables)} native DOCX table(s)")
                full_text += "\n\n[TABLES FROM DOCUMENT:]\n" + table_text

            # CRITICAL FIX: Extract and OCR embedded images (for table screenshots)
            ocr_text = self._extract_and_ocr_docx_images(doc, file_path)
            if ocr_text:
                print(f"   ✅ OCR extracted text from embedded images")
                full_text += "\n\n[OCR FROM EMBEDDED IMAGES:]\n" + ocr_text

            # Now chunk the full text
            chunks = self.text_splitter.split_text(full_text)

            for idx, chunk in enumerate(chunks):
                # Estimate page for this chunk
                estimated_page = (idx * self.chunk_size) // 2500 + 1

                # Detect section from chunk content
                section = self._detect_section(chunk) or current_section

                metadata = {
                    "source": file_name,
                    "source_type": doc_type,
                    "page": estimated_page,  # Rough estimate
                    "section": section,
                    "chunk_index": idx,
                    "total_chunks": len(chunks)
                }

                documents.append(LangchainDocument(
                    page_content=chunk,
                    metadata=metadata
                ))

            print(f"   ✅ Extracted {len(documents)} chunks from document")

        except Exception as e:
            print(f"   ❌ Error parsing DOCX: {e}")

        return documents

    def _extract_tables_from_docx(self, doc) -> str:
        """
        Extract native DOCX tables and format them as text.

        Args:
            doc: python-docx Document object

        Returns:
            Formatted string with all table content
        """
        if not doc.tables:
            return ""

        try:
            all_tables_text = []

            for table_idx, table in enumerate(doc.tables, 1):
                table_text = [f"\n--- Table {table_idx} ---"]

                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_cells.append(cell_text)

                    # Join cells with | delimiter to preserve structure
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        table_text.append(row_text)

                all_tables_text.append("\n".join(table_text))

            return "\n\n".join(all_tables_text)

        except Exception as e:
            print(f"   ⚠️ Warning: Could not extract DOCX tables: {e}")
            return ""

    def _extract_and_ocr_docx_images(self, doc, file_path: str) -> str:
        """
        Extract embedded images from DOCX and run OCR to capture table screenshots.

        Args:
            doc: python-docx Document object
            file_path: Path to DOCX file (needed to extract images)

        Returns:
            Combined OCR text from all embedded images
        """
        try:
            import zipfile
            from io import BytesIO

            ocr_results = []

            # DOCX is a ZIP file - extract images from it
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # List all files in the DOCX
                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                if not image_files:
                    return ""

                print(f"   🖼️  Found {len(image_files)} embedded image(s), running OCR...")

                for img_idx, img_file in enumerate(image_files, 1):
                    try:
                        # Extract image data
                        image_data = docx_zip.read(img_file)

                        # Convert to PIL Image
                        image = Image.open(BytesIO(image_data))

                        # Run OCR
                        ocr_text = pytesseract.image_to_string(image)

                        # Only keep if we got meaningful text (more than just noise)
                        if ocr_text and len(ocr_text.strip()) > 20:
                            ocr_results.append(f"\n[Embedded Image {img_idx} - {Path(img_file).name}]:\n{ocr_text.strip()}")
                            print(f"      ✓ Image {img_idx}: {len(ocr_text.strip())} characters extracted")

                    except Exception as e:
                        # Skip problematic images but don't fail
                        print(f"      ⚠️ Could not OCR image {img_idx}: {e}")
                        continue

            if ocr_results:
                return "\n\n".join(ocr_results)

            return ""

        except Exception as e:
            print(f"   ⚠️ Warning: Could not extract/OCR DOCX images: {e}")
            return ""

    def _extract_tables_from_page(self, page) -> str:
        """
        Extract table content from a PDF page using text positioning analysis.
        This helps capture table data that basic get_text() might miss.

        Args:
            page: PyMuPDF page object

        Returns:
            Formatted string with table content
        """
        try:
            # Get text with position information
            blocks = page.get_text("dict")["blocks"]

            table_text = []
            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        # Extract text from this line
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "") + " "

                        if line_text.strip():
                            table_text.append(line_text.strip())

            # Reconstruct as structured text
            if table_text:
                return "\n".join(table_text)

            return ""

        except Exception as e:
            # If table extraction fails, return empty string
            # (don't let this break the whole ingestion)
            return ""

    def _extract_and_ocr_images(self, page, page_num: int) -> str:
        """
        Extract images from PDF page and run OCR to capture text from table images.

        Args:
            page: PyMuPDF page object
            page_num: Page number for logging

        Returns:
            Combined OCR text from all images on the page
        """
        try:
            ocr_results = []
            image_list = page.get_images(full=True)

            if not image_list:
                return ""

            for img_index, img in enumerate(image_list):
                try:
                    # Get image reference
                    xref = img[0]
                    base_image = page.parent.extract_image(xref)
                    image_bytes = base_image["image"]

                    # Convert to PIL Image
                    image = Image.open(io.BytesIO(image_bytes))

                    # Run OCR
                    ocr_text = pytesseract.image_to_string(image)

                    # Only keep if we got meaningful text
                    if ocr_text and len(ocr_text.strip()) > 10:
                        ocr_results.append(f"[Image {img_index + 1}]:\n{ocr_text.strip()}")

                except Exception as e:
                    # Skip problematic images but don't fail the whole page
                    continue

            if ocr_results:
                return "\n\n".join(ocr_results)

            return ""

        except Exception as e:
            # If OCR fails entirely, return empty string
            # (don't let this break the whole ingestion)
            return ""

    def _detect_section(self, text: str) -> Optional[str]:
        """
        Detect section headers from text using common patterns.
        """
        # Common grant section patterns
        section_patterns = [
            r'^(Introduction|Background|Project Description|Objectives|Goals)',
            r'^(Methodology|Methods|Approach|Design)',
            r'^(Evaluation|Assessment|Metrics)',
            r'^(Team|Personnel|Key Personnel)',
            r'^(Budget|Costs|Financial)',
            r'^(Timeline|Schedule|Milestones)',
            r'^(Impact|Broader Impacts|Significance)',
            r'^(References|Bibliography|Citations)',
            r'^(Appendix|Appendices|Supplementary)',
            r'^(Eligibility|Requirements|Criteria)',
            r'^\d+\.\s+([A-Z][^.]+)',  # Numbered sections like "1. Introduction"
            r'^([A-Z][^.]+):',  # Sections like "Background:"
        ]

        for pattern in section_patterns:
            match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
            if match:
                return match.group(1) if match.lastindex else match.group(0)

        return None

    def _is_header(self, paragraph) -> bool:
        """
        Check if a DOCX paragraph is likely a header.
        """
        # Check for heading styles
        if paragraph.style.name.startswith('Heading'):
            return True

        # Check for bold, larger font
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold and first_run.font.size:
                return True

        return False


class DocumentIngester:
    """
    Main ingester that coordinates parsing and vector store creation.
    """

    def __init__(self):
        self.chunker = SmartChunker()
        self.embeddings = OpenAIEmbeddings(
            model=os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")
        )
        self.vector_stores = {}  # Separate stores for different document types

    def ingest_document(
        self,
        file_path: str,
        doc_type: str,
        create_vector_store: bool = True
    ) -> Tuple[List[LangchainDocument], Optional[FAISS]]:
        """
        Ingest a single document and optionally create a vector store.

        Args:
            file_path: Path to document
            doc_type: Type identifier (solicitation, narrative, etc.)
            create_vector_store: Whether to create FAISS index

        Returns:
            Tuple of (documents, vector_store)
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Route to appropriate parser
        if file_path.suffix.lower() == '.pdf':
            documents = self.chunker.parse_pdf(str(file_path), doc_type)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            documents = self.chunker.parse_docx(str(file_path), doc_type)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        if not documents:
            print(f"⚠️  Warning: No content extracted from {file_path.name}")
            return [], None

        # Create vector store
        vector_store = None
        if create_vector_store:
            print(f"🔮 Creating FAISS index for {doc_type}...")
            vector_store = FAISS.from_documents(documents, self.embeddings)
            self.vector_stores[doc_type] = vector_store
            print(f"   ✅ Indexed {len(documents)} chunks")

        return documents, vector_store

    def ingest_multiple(
        self,
        file_paths: Dict[str, str],
        combine_stores: bool = False
    ) -> Dict[str, FAISS]:
        """
        Ingest multiple documents and create vector stores.

        Args:
            file_paths: Dict mapping doc_type to file_path
                       e.g., {"solicitation": "path/to/rfp.pdf", "narrative": "path/to/draft.docx"}
            combine_stores: Whether to combine all documents into single store

        Returns:
            Dict mapping doc_type to vector_store
        """
        all_documents = []

        for doc_type, file_path in file_paths.items():
            print(f"\n{'='*60}")
            print(f"Processing: {doc_type.upper()}")
            print('='*60)

            docs, store = self.ingest_document(file_path, doc_type, create_vector_store=not combine_stores)
            all_documents.extend(docs)

        # If combining, create one unified store
        if combine_stores and all_documents:
            print(f"\n🔮 Creating unified FAISS index...")
            combined_store = FAISS.from_documents(all_documents, self.embeddings)
            self.vector_stores["combined"] = combined_store
            print(f"   ✅ Indexed {len(all_documents)} total chunks")

        return self.vector_stores

    def save_vector_store(self, doc_type: str, save_path: str):
        """Save a vector store to disk."""
        if doc_type not in self.vector_stores:
            raise ValueError(f"No vector store found for: {doc_type}")

        Path(save_path).parent.mkdir(parents=True, exist_ok=True)
        self.vector_stores[doc_type].save_local(save_path)
        print(f"💾 Saved {doc_type} vector store to: {save_path}")

    def load_vector_store(self, doc_type: str, load_path: str):
        """Load a vector store from disk."""
        if not Path(load_path).exists():
            raise FileNotFoundError(f"Vector store not found: {load_path}")

        self.vector_stores[doc_type] = FAISS.load_local(
            load_path,
            self.embeddings,
            allow_dangerous_deserialization=True
        )
        print(f"📂 Loaded {doc_type} vector store from: {load_path}")

    def search(
        self,
        query: str,
        doc_type: str = "combined",
        k: int = 5,
        score_threshold: float = 0.0
    ) -> List[Tuple[LangchainDocument, float]]:
        """
        Semantic search across indexed documents.

        Args:
            query: Search query
            doc_type: Which vector store to search
            k: Number of results to return
            score_threshold: Minimum relevance score

        Returns:
            List of (document, score) tuples with citations
        """
        if doc_type not in self.vector_stores:
            raise ValueError(f"No vector store found for: {doc_type}. Available: {list(self.vector_stores.keys())}")

        results = self.vector_stores[doc_type].similarity_search_with_score(query, k=k)

        # Filter by score threshold
        filtered_results = [(doc, score) for doc, score in results if score >= score_threshold]

        return filtered_results

    def format_citation(self, doc: LangchainDocument) -> str:
        """
        Format a document chunk as a citation.

        Returns: "[Source: filename, Page X, Section Y]"
        """
        meta = doc.metadata
        citation = f"[{meta.get('source', 'Unknown')}, Page {meta.get('page', '?')}"

        if meta.get('section'):
            citation += f", Section: {meta['section']}"

        citation += "]"
        return citation


# CLI Interface for testing
if __name__ == "__main__":
    import sys

    print("=" * 60)
    print("📚 GrantScout: Smart Document Ingester")
    print("=" * 60)

    if len(sys.argv) < 3:
        print("\nUsage: python ingester.py <doc_type> <file_path>")
        print("\nExample:")
        print("  python ingester.py solicitation data/rfp.pdf")
        print("  python ingester.py narrative data/draft.docx")
        sys.exit(1)

    doc_type = sys.argv[1]
    file_path = sys.argv[2]

    ingester = DocumentIngester()
    docs, store = ingester.ingest_document(file_path, doc_type)

    print(f"\n✅ Successfully ingested {len(docs)} chunks")

    # Test search if vector store was created
    if store:
        print("\n" + "=" * 60)
        print("Testing semantic search...")
        test_query = input("Enter a test query: ")
        results = ingester.search(test_query, doc_type, k=3)

        print(f"\nTop 3 results for '{test_query}':\n")
        for idx, (doc, score) in enumerate(results, 1):
            print(f"{idx}. {ingester.format_citation(doc)}")
            print(f"   Score: {score:.4f}")
            print(f"   Content: {doc.page_content[:200]}...")
            print()
