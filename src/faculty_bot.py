"""
Polymorphic Faculty Fetcher (The Scout)
========================================
Builds standardized competency profiles from mixed data sources:
- HTML faculty profile pages
- Direct PDF CV/Vita links

Output: Unified Team_Competency.json with citations
"""

import os
import re
import json
import time
import requests
from typing import Dict, List, Optional, Union
from pathlib import Path
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from openai import OpenAI
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class FacultyBot:
    """Polymorphic scraper that handles both HTML profiles and PDF vitas."""

    def __init__(self):
        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        self.max_retries = int(os.getenv("MAX_RETRIES", 3))
        self.timeout = int(os.getenv("REQUEST_TIMEOUT", 30))

    def fetch_faculty_data(self, path_or_url: str) -> Optional[Dict]:
        """
        Main polymorphic router: Detects file/URL type and delegates to appropriate parser.

        Args:
            path_or_url: Faculty profile URL, local file path (HTML/PDF/DOCX)

        Returns:
            Standardized competency profile dict or None if failed
        """
        try:
            print(f"🔍 Analyzing: {path_or_url}")

            # Check if it's a local file path
            if os.path.exists(path_or_url):
                print("   📁 Detected: Local file")
                return self._parse_local_file(path_or_url)

            # Otherwise, treat as URL
            # First, make a HEAD request to detect content type
            head_response = requests.head(path_or_url, allow_redirects=True, timeout=self.timeout)
            content_type = head_response.headers.get('Content-Type', '').lower()

            print(f"   Content-Type: {content_type}")

            # Branch based on content type
            if 'application/pdf' in content_type or path_or_url.lower().endswith('.pdf'):
                print("   📄 Detected: PDF Vita")
                return self._parse_pdf_vita(path_or_url)
            else:
                print("   🌐 Detected: HTML Profile")
                return self._scrape_html_profile(path_or_url)

        except requests.RequestException as e:
            print(f"❌ Error fetching {path_or_url}: {e}")
            return None
        except Exception as e:
            print(f"❌ Error processing {path_or_url}: {e}")
            return None

    def _parse_local_file(self, file_path: str) -> Optional[Dict]:
        """
        Parse local file (HTML, PDF, or DOCX).

        Args:
            file_path: Path to local file

        Returns:
            Standardized competency profile dict or None if failed
        """
        file_path_lower = file_path.lower()

        if file_path_lower.endswith('.html') or file_path_lower.endswith('.htm'):
            print("   🌐 Parsing local HTML file")
            return self._parse_local_html(file_path)
        elif file_path_lower.endswith('.pdf'):
            print("   📄 Parsing local PDF file")
            return self._parse_local_pdf(file_path)
        elif file_path_lower.endswith('.docx'):
            print("   📝 Parsing local DOCX biosketch")
            return self._parse_local_docx(file_path)
        else:
            print(f"   ❌ Unsupported file type: {file_path}")
            return None

    def _parse_local_html(self, file_path: str) -> Optional[Dict]:
        """Parse local HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Use LLM extraction for more robust parsing
            full_text = soup.get_text(separator='\n', strip=True)

            # Truncate if too long
            if len(full_text) > 8000:
                full_text = full_text[:8000] + "\n\n[... truncated for length ...]"

            profile = self._llm_extract_from_html(full_text, file_path)
            return profile

        except Exception as e:
            print(f"   ❌ Error parsing local HTML: {e}")
            return None

    def _parse_local_pdf(self, file_path: str) -> Optional[Dict]:
        """Parse local PDF file."""
        try:
            # Extract text with PyMuPDF
            raw_text = ""
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc, 1):
                    text = page.get_text()
                    raw_text += f"\n--- Page {page_num} ---\n{text}"

            # Truncate if too long
            if len(raw_text) > 8000:
                raw_text = raw_text[:8000] + "\n\n[... truncated for length ...]"

            # Use LLM to structure the raw text
            profile = self._llm_extract_from_vita(raw_text, file_path)
            return profile

        except Exception as e:
            print(f"   ❌ Error parsing local PDF: {e}")
            return None

    def _parse_local_docx(self, file_path: str) -> Optional[Dict]:
        """
        Parse local DOCX biosketch file (e.g., NSF Biographical Sketch).
        Similar to AgriGrantCoach's approach.
        """
        try:
            from docx import Document

            doc = Document(file_path)
            raw_text = ""

            # Extract all paragraphs
            for para in doc.paragraphs:
                raw_text += para.text + "\n"

            # Also extract text from tables (biosketches often have tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        raw_text += cell.text + "\n"

            # Truncate if too long
            if len(raw_text) > 8000:
                raw_text = raw_text[:8000] + "\n\n[... truncated for length ...]"

            # Use LLM to structure the biosketch
            profile = self._llm_extract_from_biosketch(raw_text, file_path)
            return profile

        except Exception as e:
            print(f"   ❌ Error parsing DOCX biosketch: {e}")
            return None

    def _llm_extract_from_html(self, html_text: str, source_path: str) -> Dict:
        """Use LLM to extract structured data from HTML text."""
        prompt = f"""You are analyzing a faculty member's HTML profile page. Extract the following information into a structured JSON format:

**Required Fields:**
- name: Full name of the faculty member
- email: Email address (if available)
- title: Current position/title
- department: Department or unit
- institution: University/institution name
- research_interests: List of research areas/interests (array of strings)
- education: Array of degree objects with fields: degree, field, institution, year
- publications: Array of recent publication titles (top 5-10 most relevant)
- expertise_keywords: List of technical skills, methodologies, or domain expertise (array of strings)
- notable_achievements: Awards, grants, or significant contributions (array of strings)
- grants: Array of recent grants with fields: title, funder, amount, year (if available)

**HTML Profile Text:**
{html_text}

Respond ONLY with valid JSON. If a field is not found, use an empty array [] or "Not specified".
"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a precise data extraction assistant. Always respond with valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )

            extracted_data = json.loads(response.choices[0].message.content)

            # Add metadata
            extracted_data["source_url"] = source_path
            extracted_data["source_type"] = "LOCAL_HTML"

            return extracted_data

        except Exception as e:
            print(f"❌ LLM extraction from HTML failed: {e}")
            return self._empty_profile(source_path, "LOCAL_HTML")

    def _llm_extract_from_biosketch(self, biosketch_text: str, source_path: str) -> Dict:
        """
        Use LLM to extract structured data from biosketch/CV text.
        Enhanced for NSF-style biographical sketches.
        """
        prompt = f"""You are analyzing a faculty member's Biographical Sketch (biosketch). This is typically an NSF-style document. Extract the following information into a structured JSON format:

**Required Fields:**
- name: Full name of the faculty member
- email: Email address (if available)
- title: Current position/title
- department: Department or unit
- institution: University/institution name
- research_interests: List of research areas/interests (array of strings)
- education: Array of degree objects with fields: degree, field, institution, year
- publications: Array of recent publication titles with years (top 10-15)
- expertise_keywords: List of technical skills, methodologies, or domain expertise (array of strings)
- notable_achievements: Awards, grants, or significant contributions (array of strings)
- grants: Array of funded grants with fields: title, funder, role, amount (if specified), year
- synergistic_activities: Key service, outreach, or professional activities
- active_projects: Current research projects

**Biosketch Text:**
{biosketch_text}

Respond ONLY with valid JSON. If a field is not found, use an empty array [] or "Not specified".
"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a precise data extraction assistant specialized in academic biosketches. Always respond with valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )

            extracted_data = json.loads(response.choices[0].message.content)

            # Add metadata
            extracted_data["source_url"] = source_path
            extracted_data["source_type"] = "DOCX_BIOSKETCH"

            return extracted_data

        except Exception as e:
            print(f"❌ LLM extraction from biosketch failed: {e}")
            return self._empty_profile(source_path, "DOCX_BIOSKETCH")

    def _parse_pdf_vita(self, url: str) -> Optional[Dict]:
        """
        Parse PDF CV/Vita using PyMuPDF + LLM extraction.

        Strategy:
        1. Download PDF and extract raw text
        2. Use LLM to structure the text into our standard format
        3. Return standardized profile
        """
        try:
            # Download PDF
            response = requests.get(url, timeout=self.timeout)
            response.raise_for_status()

            # Save temporarily
            temp_path = Path("temp_uploads") / f"temp_vita_{int(time.time())}.pdf"
            temp_path.parent.mkdir(exist_ok=True)
            temp_path.write_bytes(response.content)

            # Extract text with PyMuPDF
            raw_text = ""
            with fitz.open(temp_path) as doc:
                for page_num, page in enumerate(doc, 1):
                    text = page.get_text()
                    raw_text += f"\n--- Page {page_num} ---\n{text}"

            # Clean up temp file
            temp_path.unlink()

            # Truncate if too long (keep first 8000 chars to stay under token limits)
            if len(raw_text) > 8000:
                raw_text = raw_text[:8000] + "\n\n[... truncated for length ...]"

            # Use LLM to structure the raw text
            profile = self._llm_extract_from_vita(raw_text, url)
            return profile

        except Exception as e:
            print(f"❌ Error parsing PDF vita: {e}")
            return None

    def _llm_extract_from_vita(self, raw_text: str, source_url: str) -> Dict:
        """
        Use LLM to extract structured data from raw PDF text.
        """
        prompt = f"""You are analyzing a faculty member's CV/Vita. Extract the following information into a structured JSON format:

**Required Fields:**
- name: Full name of the faculty member
- email: Email address (if available)
- title: Current position/title
- department: Department or unit
- institution: University/institution name
- research_interests: List of research areas/interests (array of strings)
- education: Array of degree objects with fields: degree, field, institution, year
- publications: Array of recent publication titles (top 5-10 most relevant)
- expertise_keywords: List of technical skills, methodologies, or domain expertise (array of strings)
- notable_achievements: Awards, grants, or significant contributions (array of strings)

**Raw CV Text:**
{raw_text}

Respond ONLY with valid JSON. If a field is not found, use an empty array [] or "Not specified".
"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a precise data extraction assistant. Always respond with valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )

            extracted_data = json.loads(response.choices[0].message.content)

            # Add metadata
            extracted_data["source_url"] = source_url
            extracted_data["source_type"] = "PDF_VITA"

            return extracted_data

        except Exception as e:
            print(f"❌ LLM extraction failed: {e}")
            return self._empty_profile(source_url, "PDF_VITA")

    def _scrape_html_profile(self, url: str) -> Optional[Dict]:
        """
        Scrape HTML faculty profile page using BeautifulSoup + Regex.

        Strategy:
        1. Fetch HTML content
        2. Use BeautifulSoup to parse structure
        3. Apply regex patterns to extract key sections
        4. Return standardized profile
        """
        try:
            response = requests.get(url, timeout=self.timeout)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')

            # Extract name (common patterns)
            name = self._extract_name(soup)

            # Extract email
            email = self._extract_email(soup)

            # Extract title/position
            title = self._extract_title(soup)

            # Extract department
            department = self._extract_department(soup)

            # Extract research interests
            research_interests = self._extract_research_interests(soup)

            # Extract education
            education = self._extract_education(soup)

            # Extract publications
            publications = self._extract_publications(soup)

            # Extract expertise keywords from full text
            full_text = soup.get_text()
            expertise_keywords = self._extract_expertise_keywords(full_text)

            profile = {
                "name": name,
                "email": email,
                "title": title,
                "department": department,
                "institution": self._extract_institution(soup),
                "research_interests": research_interests,
                "education": education,
                "publications": publications,
                "expertise_keywords": expertise_keywords,
                "notable_achievements": self._extract_achievements(soup),
                "source_url": url,
                "source_type": "HTML_PROFILE"
            }

            return profile

        except Exception as e:
            print(f"❌ Error scraping HTML profile: {e}")
            return None

    # HTML Extraction Helper Methods
    # ================================

    def _extract_name(self, soup: BeautifulSoup) -> str:
        """Extract faculty name from common HTML patterns."""
        # Try common selectors
        selectors = [
            'h1.profile-name',
            'h1.faculty-name',
            'div.name h1',
            'h1',
            '.profile-header h1',
            '#faculty-name'
        ]

        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                return element.get_text(strip=True)

        return "Not specified"

    def _extract_email(self, soup: BeautifulSoup) -> str:
        """Extract email using regex."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        text = soup.get_text()
        match = re.search(email_pattern, text)
        return match.group(0) if match else "Not specified"

    def _extract_title(self, soup: BeautifulSoup) -> str:
        """Extract job title/position."""
        selectors = [
            '.profile-title',
            '.faculty-title',
            'h2.title',
            '.position',
            'p.title'
        ]

        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                return element.get_text(strip=True)

        # Fallback: look for common title keywords
        text = soup.get_text()
        title_match = re.search(r'(Professor|Associate Professor|Assistant Professor|Lecturer|Instructor)', text, re.IGNORECASE)
        return title_match.group(0) if title_match else "Not specified"

    def _extract_department(self, soup: BeautifulSoup) -> str:
        """Extract department name."""
        selectors = [
            '.department',
            '.dept',
            'div.affiliation',
            'p.department'
        ]

        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                return element.get_text(strip=True)

        return "Not specified"

    def _extract_institution(self, soup: BeautifulSoup) -> str:
        """Extract institution name."""
        # Try meta tags first
        meta = soup.find('meta', {'property': 'og:site_name'})
        if meta and meta.get('content'):
            return meta['content']

        # Try title tag
        title = soup.find('title')
        if title:
            # Often format is "Name | Department | University"
            parts = title.get_text().split('|')
            if len(parts) > 1:
                return parts[-1].strip()

        return "Not specified"

    def _extract_research_interests(self, soup: BeautifulSoup) -> List[str]:
        """Extract research interests/areas."""
        interests = []

        # Look for sections with research interests
        headers = soup.find_all(['h2', 'h3', 'h4'])
        for header in headers:
            header_text = header.get_text(strip=True).lower()
            if any(keyword in header_text for keyword in ['research interest', 'research area', 'expertise']):
                # Get the next sibling elements (likely a list or paragraph)
                next_elem = header.find_next_sibling()
                if next_elem:
                    if next_elem.name == 'ul':
                        interests.extend([li.get_text(strip=True) for li in next_elem.find_all('li')])
                    else:
                        # Split by common delimiters
                        text = next_elem.get_text()
                        interests.extend([i.strip() for i in re.split(r'[,;•]', text) if i.strip()])

        return interests[:10] if interests else []

    def _extract_education(self, soup: BeautifulSoup) -> List[Dict]:
        """Extract education history."""
        education = []

        # Look for education section
        headers = soup.find_all(['h2', 'h3', 'h4'])
        for header in headers:
            header_text = header.get_text(strip=True).lower()
            if 'education' in header_text:
                next_elem = header.find_next_sibling()
                if next_elem:
                    text = next_elem.get_text()
                    # Extract degree patterns like "Ph.D. in Computer Science, MIT, 2015"
                    degree_pattern = r'(Ph\.?D\.?|M\.?S\.?|B\.?S\.?|M\.?A\.?|B\.?A\.?)[,\s]+(?:in\s+)?([^,]+),\s*([^,]+),?\s*(\d{4})?'
                    matches = re.findall(degree_pattern, text, re.IGNORECASE)
                    for match in matches:
                        education.append({
                            "degree": match[0],
                            "field": match[1].strip(),
                            "institution": match[2].strip(),
                            "year": match[3] if match[3] else "Not specified"
                        })

        return education

    def _extract_publications(self, soup: BeautifulSoup) -> List[str]:
        """Extract publication titles."""
        publications = []

        # Look for publications section
        headers = soup.find_all(['h2', 'h3', 'h4'])
        for header in headers:
            header_text = header.get_text(strip=True).lower()
            if 'publication' in header_text or 'paper' in header_text:
                # Find list items or paragraphs after this header
                next_elem = header.find_next_sibling()
                count = 0
                while next_elem and count < 10:
                    if next_elem.name == 'ul':
                        pubs = [li.get_text(strip=True) for li in next_elem.find_all('li')[:10]]
                        publications.extend(pubs)
                        break
                    elif next_elem.name in ['p', 'div']:
                        pub_text = next_elem.get_text(strip=True)
                        if len(pub_text) > 20:  # Likely a publication
                            publications.append(pub_text)
                            count += 1
                    next_elem = next_elem.find_next_sibling()

        return publications[:10]

    def _extract_expertise_keywords(self, text: str) -> List[str]:
        """Extract technical keywords from full text."""
        # Common technical terms in various domains
        keywords_pattern = r'\b(machine learning|deep learning|artificial intelligence|data science|statistics|'
        keywords_pattern += r'computer vision|natural language processing|NLP|neural networks|'
        keywords_pattern += r'python|R|java|C\+\+|matlab|tensorflow|pytorch|'
        keywords_pattern += r'algorithm|optimization|simulation|modeling|analysis|'
        keywords_pattern += r'database|SQL|cloud computing|AWS|Azure|'
        keywords_pattern += r'cybersecurity|cryptography|blockchain|IoT|'
        keywords_pattern += r'genomics|bioinformatics|proteomics|CRISPR|'
        keywords_pattern += r'climate|sustainability|renewable energy|'
        keywords_pattern += r'quantum computing|robotics|automation)\b'

        matches = re.findall(keywords_pattern, text, re.IGNORECASE)
        # Deduplicate and normalize
        unique_keywords = list(set([kw.lower() for kw in matches]))
        return unique_keywords[:15]

    def _extract_achievements(self, soup: BeautifulSoup) -> List[str]:
        """Extract awards, grants, honors."""
        achievements = []

        headers = soup.find_all(['h2', 'h3', 'h4'])
        for header in headers:
            header_text = header.get_text(strip=True).lower()
            if any(keyword in header_text for keyword in ['award', 'honor', 'grant', 'achievement']):
                next_elem = header.find_next_sibling()
                if next_elem:
                    if next_elem.name == 'ul':
                        achievements.extend([li.get_text(strip=True) for li in next_elem.find_all('li')[:5]])
                    else:
                        text = next_elem.get_text(strip=True)
                        if text:
                            achievements.append(text)

        return achievements[:5]

    def _empty_profile(self, url: str, source_type: str) -> Dict:
        """Return an empty profile structure when extraction fails."""
        return {
            "name": "Not specified",
            "email": "Not specified",
            "title": "Not specified",
            "department": "Not specified",
            "institution": "Not specified",
            "research_interests": [],
            "education": [],
            "publications": [],
            "expertise_keywords": [],
            "notable_achievements": [],
            "source_url": url,
            "source_type": source_type
        }

    def process_faculty_list(self, urls: List[str], output_path: str = None) -> Dict:
        """
        Process multiple faculty URLs and generate Team_Competency.json

        Args:
            urls: List of faculty profile URLs (mixed HTML and PDF)
            output_path: Optional path to save JSON output

        Returns:
            Team competency dictionary
        """
        print(f"\n🚀 Processing {len(urls)} faculty members...")

        team_data = {
            "team_members": [],
            "metadata": {
                "total_members": len(urls),
                "successful_fetches": 0,
                "failed_fetches": 0,
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
            }
        }

        for idx, url in enumerate(urls, 1):
            print(f"\n[{idx}/{len(urls)}] Processing: {url}")
            profile = self.fetch_faculty_data(url)

            if profile:
                team_data["team_members"].append(profile)
                team_data["metadata"]["successful_fetches"] += 1
                print(f"   ✅ Success: {profile.get('name', 'Unknown')}")
            else:
                team_data["metadata"]["failed_fetches"] += 1
                print(f"   ❌ Failed to fetch profile")

        # Save to file if path provided
        if output_path:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(team_data, f, indent=2, ensure_ascii=False)
            print(f"\n💾 Saved team competency data to: {output_path}")

        return team_data


# CLI Interface for testing
if __name__ == "__main__":
    import sys

    print("=" * 60)
    print("🔬 GrantScout: Polymorphic Faculty Fetcher")
    print("=" * 60)

    if len(sys.argv) < 2:
        print("\nUsage: python faculty_bot.py <url1> <url2> ...")
        print("\nExample:")
        print("  python faculty_bot.py https://example.edu/~faculty/profile.html")
        print("  python faculty_bot.py https://example.edu/cv.pdf")
        sys.exit(1)

    urls = sys.argv[1:]
    bot = FacultyBot()
    team_data = bot.process_faculty_list(urls, output_path="data/Team_Competency.json")

    print("\n" + "=" * 60)
    print(f"✅ Processed {team_data['metadata']['successful_fetches']}/{team_data['metadata']['total_members']} profiles successfully")
    print("=" * 60)
